import os
import uuid
import logging
from docx import Document

logger = logging.getLogger(__name__)

class DOCXResumeRenderer:
    """Compiles a Structured Resume Specification into a clean, downloadable DOCX file."""

    @staticmethod
    def render(spec_data: dict, output_dir: str = None) -> str:
        """Create a DOCX file on disk and return the path."""
        if not output_dir:
            output_dir = os.path.join(os.getcwd(), "media", "resumes")
        os.makedirs(output_dir, exist_ok=True)

        file_id = str(uuid.uuid4())[:8]
        output_path = os.path.join(output_dir, f"resume_{file_id}.docx")

        try:
            doc = Document()
            header = spec_data.get("header", {})
            
            # Header
            doc.add_heading(header.get("full_name", "Shivam Singh"), 0)
            p = doc.add_paragraph()
            p.add_run(header.get("headline", "")).bold = True
            
            contact_info = []
            if header.get("email"): contact_info.append(header["email"])
            if header.get("phone"): contact_info.append(header["phone"])
            if header.get("location"): contact_info.append(header["location"])
            
            doc.add_paragraph(" | ".join(contact_info))
            
            # Summary
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(spec_data.get("summary", ""))
            
            # Skills
            doc.add_heading("Technical Skills", level=1)
            for g in spec_data.get("skills_groups", []):
                dp = doc.add_paragraph()
                dp.add_run(f"{g.get('category', 'Skills')}: ").bold = True
                dp.add_run(", ".join(g.get("skills", [])))

            # Experience
            doc.add_heading("Professional Experience", level=1)
            for exp in spec_data.get("experiences", []):
                dp = doc.add_paragraph()
                dp.add_run(f"{exp.get('role')} at {exp.get('company')}").bold = True
                dp.add_run(f" ({exp.get('start_date')} - {exp.get('end_date')})")
                for bullet in exp.get("bullet_points", []):
                    doc.add_paragraph(bullet, style='List Bullet')
                    
            # Projects
            doc.add_heading("Key Projects", level=1)
            for proj in spec_data.get("projects", []):
                doc.add_paragraph(proj.get("title", "Project")).bold = True
                for bullet in proj.get("bullet_points", []):
                    doc.add_paragraph(bullet, style='List Bullet')

            doc.save(output_path)
            return output_path
        except Exception as e:
            logger.error(f"Failed to generate DOCX resume: {e}")
            raise
